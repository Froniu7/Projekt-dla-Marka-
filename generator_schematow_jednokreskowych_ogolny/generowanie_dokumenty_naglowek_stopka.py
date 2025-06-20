
from docx import Document
from docx.shared import Inches
import os

# Ścieżki do grafik
naglowek_obrazek = "naglowek.png"
stopka_obrazek = "stopka.png"

# Tworzenie dokumentu
doc = Document()

# Marginesy
cm = 0.3937  # 1 cm
section = doc.sections[0]
section.top_margin = Inches(cm)
section.left_margin = Inches(cm)
section.right_margin = Inches(cm)
section.bottom_margin = Inches(0.19685)       # 0.5 cm
section.footer_distance = Inches(0.1)
# Szerokość użytkowa
page_width = section.page_width.inches
usable_width = page_width - section.left_margin.inches - section.right_margin.inches

# === Nagłówek ===
header = section.header
for p in header.paragraphs:
    p.clear()
header_paragraph = header.add_paragraph()
header_paragraph.add_run().add_picture(naglowek_obrazek, width=Inches(usable_width))

# === Stopka ===
footer = section.footer
for p in footer.paragraphs:
    p.clear()

# Obrazek w stopce
footer_paragraph_img = footer.add_paragraph()
footer_paragraph_img.alignment = 1  # Wyśrodkowanie
footer_paragraph_img.add_run().add_picture(stopka_obrazek, width=Inches(usable_width))

# Tekst pod grafiką
footer_paragraph_text = footer.add_paragraph("Strona 1")
footer_paragraph_text.alignment = 1  # Wyśrodkowanie

# Treść główna
doc.add_paragraph("To jest treść dokumentu między grafiką w nagłówku i w stopce.")

# Zapisz dokument
doc.save("8.docx")
print("✔ 8.docx z tekstem pod grafiką w stopce")







