from PIL import Image




def generuj_grafike_z_tekstem(tekst, szerokosc, wysokosc, rozmiar_fonta, nazwa_pliku):
    from PIL import Image, ImageDraw, ImageFont
    # Tworzymy pusty biały obraz
    obraz = Image.new('RGBA', (szerokosc, wysokosc), color=(0,0,0,0))

    # Obiekt do rysowania
    draw = ImageDraw.Draw(obraz)

    # Czcionka – jeśli nie masz arial.ttf, załadujemy domyślną
    try:
        czcionka = ImageFont.truetype("arial.ttf", rozmiar_fonta)
    except IOError:
        czcionka = ImageFont.load_default()

    # Oblicz wymiary tekstu za pomocą textbbox
    bbox = draw.textbbox((0, 0), f"{tekst}", font=czcionka)
    tekst_szerokosc = bbox[2] - bbox[0]
    tekst_wysokosc = bbox[3] - bbox[1]

    # Środek obrazu
    pozycja = ((szerokosc - tekst_szerokosc) // 2, (wysokosc - tekst_wysokosc) // 2)

    # Rysuj tekst
    draw.text(pozycja, f"{tekst}", fill='black', font=czcionka)

    # Zapisz
    obraz.save(nazwa_pliku)
    print(f"Grafika zapisana jako {nazwa_pliku}")

def generuj_grafiki_do_schematu(cable, control, przekladnik, zabezpieczenie):
    generuj_grafike_z_tekstem(tekst=cable, szerokosc=300, wysokosc=80, rozmiar_fonta=50,
                              nazwa_pliku="przewod_kablowy.png")
    generuj_grafike_z_tekstem(tekst=control, szerokosc=400, wysokosc=100, rozmiar_fonta=50,
                              nazwa_pliku="przewod_sterowniczy.png")
    generuj_grafike_z_tekstem(tekst=przekladnik, szerokosc=600, wysokosc=100, rozmiar_fonta=50,
                              nazwa_pliku="przekladnik.png")
    generuj_grafike_z_tekstem(tekst=zabezpieczenie, szerokosc=100, wysokosc=100, rozmiar_fonta=50,
                              nazwa_pliku="zabezpiecznie.png")


def generowanie_schematu_jednokreskowego(rodzaj_zab):
    print("Generowanie schematu jednokreskowego.")
    if rodzaj_zab == "typ S Cx A":
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_C/schemat_jednokreskowy_SVC_C.png")
    elif rodzaj_zab == "wkładki gG":
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_gG/schemat_jednokreskowy_SVC_gG.png")
    else:
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_brak/schemat_jednokreskowy_SVC_brak.png")
    width_schemat, height_schemat = schemat_jednokreskowy.size
    print(f"width {width_schemat}")
    print(f"height {height_schemat}")

    przewod_kablowy = Image.open("przewod_kablowy.png")
    przewod_sterowniczy = Image.open("przewod_sterowniczy.png")
    przekladnik = Image.open("przekladnik.png")
    zabezpieczenie = Image.open("zabezpiecznie.png")
    obraz_do_wygenerowania = Image.new("RGBA", (width_schemat, height_schemat))
    obraz_do_wygenerowania.paste(schemat_jednokreskowy, (0, 0))
    obraz_do_wygenerowania.paste(przewod_kablowy, (2300, 980))
    obraz_do_wygenerowania.paste(przewod_sterowniczy,(1300,1000))
    obraz_do_wygenerowania.paste(przekladnik, (1230, 100))
    obraz_do_wygenerowania.paste(zabezpieczenie,(3000,800))

    obraz_do_wygenerowania.save(f"schemat_jednokreskowy.png")
    print("Został wygenerowany shemat jednokreskowy.")

def generowanie_schematu_podlaczenia(rodzaj_zab, sterownik):
    print("Generowanie schematu jednokreskowego.")
    if (rodzaj_zab == "typ S Cx A") and (sterownik == "Novar"):
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_C/schemat_sczegolowy_podpiecia_Novar.png")
    elif (rodzaj_zab == "typ S Cx A") and (sterownik == "Tense"):
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_C/schemat_sczegolowy_podpiecia_Active.png")

    elif (rodzaj_zab == "wkładki gG") and (sterownik == "Novar"):
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_gG/schemat_sczegolowy_podpiecia_Novar.png")
    elif (rodzaj_zab == "wkładki gG") and (sterownik == "Tense"):
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_gG/schemat_sczegolowy_podpiecia_Active.png.png")

    elif (rodzaj_zab == "brak zab. głównego zewnętrznego") and (sterownik == "Novar"):
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_brak/schemat_sczegolowy_podpiecia_Novar.png")
    else:
        schemat_jednokreskowy = Image.open("schematy_zabezpiecznie_brak/schemat_sczegolowy_podpiecia_Active.png")
    width_schemat, height_schemat = schemat_jednokreskowy.size
    print(f"wymiary schemat ogolny - szerokosc : {width_schemat}")
    print(f"wymiary schemat ogolny - wysokosc : {height_schemat}")
    przewod_kablowy = Image.open("przewod_kablowy.png")
    przewod_sterowniczy = Image.open("przewod_sterowniczy.png")
    przekladnik = Image.open("przekladnik.png")
    zabezpieczenie = Image.open("zabezpiecznie.png")
    obraz_do_wygenerowania = Image.new("RGBA", (width_schemat, height_schemat))
    #schemat glowny
    obraz_do_wygenerowania.paste(schemat_jednokreskowy, (0, 0))
    #Przewod - Zasilanie
    obraz_do_wygenerowania.paste(przewod_kablowy, (2100, 980))
    #Sterowniczy
    obraz_do_wygenerowania.paste(przewod_sterowniczy, (200, 600))
    #Przekladniki
    obraz_do_wygenerowania.paste(przekladnik, (610, 120))
    #Zabezpiecznie
    obraz_do_wygenerowania.paste(zabezpieczenie, (1700, 600))

    obraz_do_wygenerowania.save(f"schemat_ogolny.png")
    print("Został wygenerowany shemat ogolny.")

def generuj_plik_docx(sciezka_jednokreskowy, sciezka_ogolny, nazwa_pliku, przekladniki, zasilajacy, sterowniczy, rodzaj_zab, ampery):
    from docx import Document
    from docx.shared import Cm

    doc = Document()

    # Ustaw marginesy na 1 cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Oblicz szerokość do wykorzystania: szerokość strony - marginesy
    section = doc.sections[0]
    szerokosc_strony = section.page_width  # w EMU
    lewy_margin = section.left_margin
    prawy_margin = section.right_margin
    szerokosc_do_wykorzystania = szerokosc_strony - lewy_margin - prawy_margin

    doc.add_paragraph("Schemat jendnokreskowy - podpięcia urządzenia do kompensacji:")

    doc.add_picture(sciezka_jednokreskowy, width=szerokosc_do_wykorzystania)
    legenda = doc.add_paragraph("Legenda:\n")
    run = legenda.add_run(f"Przekładniki : {przekladniki}").add_break()
    run = legenda.add_run(f"Przewód kablowy zasilający : {zasilajacy}").add_break()
    run = legenda.add_run(f"Przewód kablowy sterowniczy : {sterowniczy}").add_break()
    if rodzaj_zab == "typ S Cx A":
        run = legenda.add_run(f"Zabezpieczenie typu 'S' C 3-polowe {ampery}").add_break()
    if rodzaj_zab == "wkładki gG":
        run = legenda.add_run(f"Zabezpieczenie wkładki Gg {ampery}").add_break()
    doc.add_page_break()
    doc.add_paragraph("Schemat szczegółowy - podpięcia urządzenia do kompensacji:")
    doc.add_picture(sciezka_ogolny, width=szerokosc_do_wykorzystania)
    legenda = doc.add_paragraph("Legenda:\n")
    run = legenda.add_run(f"Przekładniki : {przekladniki}").add_break()
    run = legenda.add_run(f"Przewód kablowy zasilający : {zasilajacy}").add_break()
    run = legenda.add_run(f"Przewód kablowy sterowniczy : {sterowniczy}").add_break()
    if rodzaj_zab == "typ S Cx A":
        run = legenda.add_run(f"Zabezpieczenie typu 'S' C 3-polowe {ampery}").add_break()
    if rodzaj_zab == "wkładki gG":
        run = legenda.add_run(f"Zabezpieczenie wkładki Gg {ampery}").add_break()

    '''
    legenda = doc.add_paragraph("Legenda:\n")
    run = legenda.add_run(f"Q1 - dławik 1 fazowy, moc {q1} kVAr").add_break()
    run = legenda.add_run(f"Q2 - dławik 1 fazowy, moc {q2} kVAr").add_break()
    run = legenda.add_run(f"Q3 - dławik 1 fazowy, moc {q3} kVAr").add_break()
    liczymy_stopnie = 3
    for i in range(12):

        if lista[i] != "":
            liczymy_stopnie += 1
            if i <= 5:
                if lista[i].startswith("-"):
                    run = legenda.add_run(f"Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr").add_break()
                    print(f" doloczono do docx opis dla dlawik 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")
                else:
                    run = legenda.add_run(f"Q{liczymy_stopnie} - kondensator 1 fazowy, moc {lista[i]} kVAr").add_break()
                    print(
                        f" doloczono do docx opis dla kondensator 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")
            if i > 5:
                if lista[i].startswith("-"):
                    run = legenda.add_run(f"Q{liczymy_stopnie} - dławik 3 fazowy, moc {lista[i]} kVAr").add_break()
                    print(f" doloczono do docx opis dla dlawik 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")
                else:
                    run = legenda.add_run(f"Q{liczymy_stopnie} - kondensator 3 fazowy, moc {lista[i]} kVAr").add_break()
                    print(
                        f" doloczono do docx opis dla kondensator 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")

    float_d1 = float(q1)
    float_d2 = float(q2)
    float_d3 = float(q3)
    suma_mocy_lacznik_tyrystorowy = (float_d1 + float_d2 + float_d3)*(-1)
    print(f"suma mocy dla łącznika tyrystorowego to : {suma_mocy_lacznik_tyrystorowy}")
    if suma_mocy_lacznik_tyrystorowy < 5:
        run = legenda.add_run("Łącznik tyrystorowy o mocy 5 kVAr" ).add_break()
    elif (suma_mocy_lacznik_tyrystorowy > 5) and (suma_mocy_lacznik_tyrystorowy <= 10):
        run = legenda.add_run("Łącznik tyrystorowy o mocy 10 kVAr").add_break()
    elif (suma_mocy_lacznik_tyrystorowy > 10) and (suma_mocy_lacznik_tyrystorowy <= 15):
        run = legenda.add_run("Łącznik tyrystorowy o mocy 15 kVAr").add_break()
    elif (suma_mocy_lacznik_tyrystorowy > 15) and (suma_mocy_lacznik_tyrystorowy <= 20):
        run = legenda.add_run("Łącznik tyrystorowy o mocy 20 kVAr").add_break()
    run = legenda.add_run("Sterownik z funkcją SVC").add_break()
    run = legenda.add_run("Wentylator 230V 0.12A, załączany termostatem KTS 011").add_break()
    '''
    print("zapisujemy plik")
    doc.save(nazwa_pliku)

#generuj_plik_docx("schemat_jednokreskowy.png","schemat_ogolny.png","dokument_schematy.docx")





