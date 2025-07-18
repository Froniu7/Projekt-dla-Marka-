

def generuj_grafike_z_tekstem(tekst, szerokosc, wysokosc, rozmiar_fonta, nazwa_pliku):
    from PIL import Image, ImageDraw, ImageFont
    # Tworzymy pusty biały obraz
    obraz = Image.new('RGBA', (szerokosc, wysokosc), color=(0,0,0,0))

    # Obiekt do rysowania
    draw = ImageDraw.Draw(obraz)

    # Czcionka – jeśli nie masz arial.ttf, załadujemy domyślną
    try:
        czcionka = ImageFont.truetype("arial.ttf", rozmiar_fonta)
    except IOError:
        czcionka = ImageFont.load_default()

    # Oblicz wymiary tekstu za pomocą textbbox
    bbox = draw.textbbox((0, 0), f"{tekst} kVAr", font=czcionka)
    tekst_szerokosc = bbox[2] - bbox[0]
    tekst_wysokosc = bbox[3] - bbox[1]

    # Środek obrazu
    pozycja = ((szerokosc - tekst_szerokosc) // 2, (wysokosc - tekst_wysokosc) // 2)

    # Rysuj tekst
    draw.text(pozycja, f"{tekst} kVAr", fill='black', font=czcionka)

    # Zapisz
    obraz.save(nazwa_pliku)
    print(f"Grafika zapisana jako {nazwa_pliku}")

#generuj_grafike_z_tekstem("4",1299, 300,100,"grafika.png")

from docx.shared import Cm

def generuj_plik_docx(sciezka_obrazu, sciezka_docx, q1, q2, q3, lista):
    from docx import Document
    from docx.shared import Inches, Cm

    doc = Document()

    # Ustaw marginesy na 1 cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Oblicz szerokość do wykorzystania: szerokość strony - marginesy
    section = doc.sections[0]
    szerokosc_strony = section.page_width  # w EMU
    lewy_margin = section.left_margin
    prawy_margin = section.right_margin
    szerokosc_do_wykorzystania = szerokosc_strony - lewy_margin - prawy_margin

    doc.add_paragraph("Schemat urządzenia SVC:")

    # Dodaj obraz o maksymalnej szerokości (EMU to jednostka używana przez python-docx)
    doc.add_picture(sciezka_obrazu, width=szerokosc_do_wykorzystania)
    legenda = doc.add_paragraph("Legenda:\n")
    run = legenda.add_run(f"Q1 - dławik 1 fazowy, moc {q1} kVAr").add_break()
    run = legenda.add_run(f"Q2 - dławik 1 fazowy, moc {q2} kVAr").add_break()
    run = legenda.add_run(f"Q3 - dławik 1 fazowy, moc {q3} kVAr").add_break()
    liczymy_stopnie = 3
    for i in range(12):

        if lista[i] != "":
            liczymy_stopnie += 1
            if i <= 5:
                if lista[i].startswith("-"):
                    run = legenda.add_run(f"Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr").add_break()
                    print(f" doloczono do docx opis dla dlawik 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")
                else:
                    run = legenda.add_run(f"Q{liczymy_stopnie} - kondensator 1 fazowy, moc {lista[i]} kVAr").add_break()
                    print(
                        f" doloczono do docx opis dla kondensator 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")
            if i > 5:
                if lista[i].startswith("-"):
                    run = legenda.add_run(f"Q{liczymy_stopnie} - dławik 3 fazowy, moc {lista[i]} kVAr").add_break()
                    print(f" doloczono do docx opis dla dlawik 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")
                else:
                    run = legenda.add_run(f"Q{liczymy_stopnie} - kondensator 3 fazowy, moc {lista[i]} kVAr").add_break()
                    print(
                        f" doloczono do docx opis dla kondensator 1 fazowy Q{liczymy_stopnie} - dławik 1 fazowy, moc {lista[i]} kVAr")

    float_d1 = float(q1)
    float_d2 = float(q2)
    float_d3 = float(q3)
    suma_mocy_lacznik_tyrystorowy = (float_d1 + float_d2 + float_d3)*(-1)
    print(f"suma mocy dla łącznika tyrystorowego to : {suma_mocy_lacznik_tyrystorowy}")
    if suma_mocy_lacznik_tyrystorowy < 5:
        run = legenda.add_run("Łącznik tyrystorowy o mocy 5 kVAr" ).add_break()
    elif (suma_mocy_lacznik_tyrystorowy > 5) and (suma_mocy_lacznik_tyrystorowy <= 10):
        run = legenda.add_run("Łącznik tyrystorowy o mocy 10 kVAr").add_break()
    elif (suma_mocy_lacznik_tyrystorowy > 10) and (suma_mocy_lacznik_tyrystorowy <= 15):
        run = legenda.add_run("Łącznik tyrystorowy o mocy 15 kVAr").add_break()
    elif (suma_mocy_lacznik_tyrystorowy > 15) and (suma_mocy_lacznik_tyrystorowy <= 20):
        run = legenda.add_run("Łącznik tyrystorowy o mocy 20 kVAr").add_break()
    run = legenda.add_run("Sterownik z funkcją SVC").add_break()
    run = legenda.add_run("Wentylator 230V 0.12A, załączany termostatem KTS 011").add_break()






    doc.save(sciezka_docx)

def konwertuj_docx_na_pdf(sciezka_docx):
    import subprocess
    import os
    sciezka_docx = os.path.abspath(sciezka_docx)
    katalog_wyjscia = os.path.dirname(sciezka_docx)

    # PEŁNA ŚCIEŻKA do LibreOffice w Windows
    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

    subprocess.run([
        libreoffice_path,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", katalog_wyjscia,
        sciezka_docx
    ], check=True)

    print(f"PDF zapisany w katalogu: {katalog_wyjscia}")


import shutil
import os
from tkinter import filedialog
from tkinter import Tk

def zapisz_plik_do_katalogu(sciezka_pliku):
    # Ukryj główne okno Tkintera
    root = Tk()
    root.withdraw()

    # Domyślna nazwa pliku
    domyslna_nazwa = os.path.basename(sciezka_pliku)

    # Okno do wyboru lokalizacji i nazwy pliku
    sciezka_docelowa = filedialog.asksaveasfilename(
        title="Zapisz plik jako",
        initialfile=domyslna_nazwa,
        defaultextension=os.path.splitext(domyslna_nazwa)[1],
        filetypes=[("Wszystkie pliki", "*.*")]
    )

    if not sciezka_docelowa:
        print("Anulowano zapis.")
        return

    try:
        shutil.copy2(sciezka_pliku, sciezka_docelowa)
        print(f"Plik zapisany do: {sciezka_docelowa}")
    except Exception as e:
        print(f"Błąd podczas zapisu pliku: {e}")



import tkinter as tk
from PIL import Image, ImageTk

def show_non_blocking_message(parent, message):
    # Wczytaj animowany GIF
    pil_image = Image.open("grafiki_GUI/animacja_loading.gif")
    frames = []

    try:
        while True:
            frame = ImageTk.PhotoImage(pil_image.copy())
            frames.append(frame)
            pil_image.seek(len(frames))  # przejdź do następnej klatki
    except EOFError:
        pass  # koniec klatek

    win = tk.Toplevel(parent)
    win.title("Ładowanie")
    win.transient(parent)

    window_width = 250
    window_height = 100
    screen_width = win.winfo_screenwidth()
    screen_height = win.winfo_screenheight()
    center_x = int(screen_width / 2 - window_width / 2)
    center_y = int(screen_height / 2 - window_height / 2)
    win.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')

    label_text = tk.Label(win, text=message)
    label_text.pack()

    label_gif = tk.Label(win)
    label_gif.pack()

    def animate(idx=0):
        label_gif.config(image=frames[idx])
        label_gif.image = frames[idx]  # 🔴 zachowaj referencję
        win.after(30, animate, (idx + 1) % len(frames))

    animate()

    return win


def close_message_window(window):
    if window.winfo_exists():
        window.destroy()

import tkinter as tk
from PIL import Image, ImageTk

class ToolTip:
    def __init__(self, widget, text, delay=1000, image_path=None):
        self.widget = widget
        self.text = text
        self.delay = delay  # opóźnienie w ms
        self.image_path = image_path  # opcjonalna ścieżka do obrazu
        self.tipwindow = None
        self.id = None
        self.image = None  # zapobiega usunięciu przez GC

        self.widget.bind("<Enter>", self.schedule)
        self.widget.bind("<Leave>", self.hide_tip)

    def schedule(self, event=None):
        self.unschedule()
        self.id = self.widget.after(self.delay, self.show_tip)

    def unschedule(self):
        if self.id:
            self.widget.after_cancel(self.id)
            self.id = None

    def show_tip(self, event=None):
        if self.tipwindow or not self.text:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 10
        self.tipwindow = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")

        frame = tk.Frame(tw, background="white", relief="solid", borderwidth=1)
        frame.pack()

        label = tk.Label(frame, text=self.text, background="white", font=("tahoma", "8", "normal"))
        label.pack(padx=5, pady=(3, 0))

        if self.image_path:
            try:
                pil_image = Image.open(self.image_path)
                #pil_image = pil_image.resize((150, 150), Image.Resampling.LANCZOS)
                self.image = ImageTk.PhotoImage(pil_image)
                image_label = tk.Label(frame, image=self.image, background="white")
                image_label.pack(padx=5, pady=3)
            except Exception as e:
                print(f"Błąd wczytywania obrazu w ToolTip: {e}")

    def hide_tip(self, event=None):
        self.unschedule()
        if self.tipwindow:
            self.tipwindow.destroy()
            self.tipwindow = None

def podglad_obrazu(program, czy_schemat_gotowy):
    import subprocess

    # Ścieżka do pliku .exe, który jest w katalogu projektu (np. w tym samym folderze co skrypt)
    exe_path = os.path.join(os.getcwd(), program)
    argument = str(czy_schemat_gotowy)
    # Otwieramy plik .exe
    #os.startfile(exe_path)
    print(exe_path)
    subprocess.Popen([exe_path, "--schemat_ready", argument])

def zamien_przecinek_na_kropke(tekst):
    if "," in tekst:
        return tekst.replace(",", ".")
    return tekst


def czy_mozna_na_float_prosty(s):
    import re
    wzorzec = r'^-?\d+(\.\d+)?$'
    return bool(re.match(wzorzec, s))

def czy_mozna_na_float_prosty_z_minus_obowiazkowo(s):
    import re
    wzorzec = r'^-\d+(\.\d+)?$'
    return bool(re.match(wzorzec, s))


